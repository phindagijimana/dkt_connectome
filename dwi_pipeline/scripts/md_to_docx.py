#!/usr/bin/env python3
"""Convert pipeline markdown documentation to .docx (basic formatting)."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH

IMAGE_RE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")
# Bold before italic, so ** is not consumed by the single-* pattern.
INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|(?<!\*)\*[^*]+\*(?!\*)|`[^`]+`)")


def add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    run = p.add_run("\n".join(lines))
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def add_rich_text(paragraph, text: str) -> None:
    """Render **bold**, *italic* and `code` as runs rather than literal markers."""
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)  # links -> label
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
        elif part.startswith("*") and part.endswith("*"):
            paragraph.add_run(part[1:-1]).italic = True
        else:
            paragraph.add_run(part)


def usable_width(doc: Document) -> Emu:
    section = doc.sections[0]
    return Emu(section.page_width - section.left_margin - section.right_margin)


def add_image(doc: Document, alt: str, src: Path, max_width: Emu) -> None:
    """Embed an image scaled to the text width, or note it if unavailable."""
    if not src.is_file():
        p = doc.add_paragraph()
        run = p.add_run(f"[missing figure: {src.name} — {alt}]")
        run.italic = True
        print(f"  WARNING: figure not found: {src}", file=sys.stderr)
        return
    doc.add_picture(str(src))
    pic = doc.inline_shapes[-1]
    if pic.width > max_width:
        pic.height = Emu(int(pic.height * max_width / pic.width))
        pic.width = max_width
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def parse_table(lines: list[str]) -> tuple[list[str], list[list[str]]]:
    rows = []
    for line in lines:
        if not line.strip().startswith("|"):
            break
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) < 2:
        return [], []
    header = rows[0]
    body = [r for r in rows[2:] if any(c.strip() for c in r)]  # skip separator
    return header, body


def md_to_docx(md_path: Path, docx_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    max_width = usable_width(doc)
    md_dir = md_path.resolve().parent

    i = 0
    in_code = False
    code_buf: list[str] = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                add_code_block(doc, code_buf)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if line.strip() == "---":
            doc.add_paragraph("")
            i += 1
            continue

        img = IMAGE_RE.match(line.strip())
        if img:
            add_image(doc, img.group(1), (md_dir / img.group(2)).resolve(), max_width)
            i += 1
            continue

        heading = re.match(r"^(#{1,4})\s+(.*)$", line)
        if heading:
            # Markers such as *(Beginner)* are part of the title text, not emphasis
            # Word can render inside a heading, so drop them.
            title = re.sub(r"[*`]", "", heading.group(2)).strip()
            doc.add_heading(title, level=len(heading.group(1)))
            i += 1
            continue

        if line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            header, body = parse_table(table_lines)
            if header:
                tbl = doc.add_table(rows=1 + len(body), cols=len(header))
                tbl.style = "Table Grid"
                for j, h in enumerate(header):
                    tbl.rows[0].cells[j].text = h
                for ri, row in enumerate(body):
                    for j, cell in enumerate(row):
                        if j < len(tbl.rows[ri + 1].cells):
                            tbl.rows[ri + 1].cells[j].text = cell
                doc.add_paragraph("")
            continue

        if line.strip().startswith("- "):
            add_rich_text(doc.add_paragraph(style="List Bullet"), line.strip()[2:])
            i += 1
            continue

        if re.match(r"^\d+\.\s", line.strip()):
            add_rich_text(
                doc.add_paragraph(style="List Number"),
                re.sub(r"^\d+\.\s", "", line.strip()),
            )
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        add_rich_text(doc.add_paragraph(), line)
        i += 1

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print(f"Wrote {docx_path}")


def main() -> None:
    if len(sys.argv) != 3:
        print(f"Usage: {sys.argv[0]} input.md output.docx", file=sys.stderr)
        sys.exit(1)
    md_to_docx(Path(sys.argv[1]), Path(sys.argv[2]))


if __name__ == "__main__":
    main()
